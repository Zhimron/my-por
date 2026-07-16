"""Generate Shimron Guray's ATS-friendly resume as PDF and DOCX."""

from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
PDF_PATH = PUBLIC / "resume.pdf"
DOCX_PATH = PUBLIC / "shimron-guray-ats-resume.docx"
ORIGINAL_PATH = PUBLIC / "resume-original.pdf"

NAME = "SHIMRON M. GURAY"
TITLE = "PROGRAMMING SPECIALIST | FULL-STACK DEVELOPER"
CONTACT = (
    "San Jose, Plaridel, Bulacan, Philippines | +63 985 490 2174 | "
    "gurayshimron@gmail.com"
)
LINKS = (
    "LinkedIn: linkedin.com/in/shimron-guray-a234472b8 | "
    "GitHub: github.com/Zhimron | Portfolio: zhimron.github.io/my-por"
)

SUMMARY = (
    "Programming Specialist and Full-Stack Developer with 3+ years of experience "
    "building web, desktop, automation, and computer vision applications. Experienced "
    "in backend development, REST API design, "
    "database integration, responsive frontend development, and legacy system "
    "modernization. Works across React, TypeScript, JavaScript, PHP, Python, Node.js, "
    "Laravel, MySQL, MongoDB, Microsoft SQL Server, and VB.NET."
)

SKILLS = [
    ("Programming Languages", "TypeScript, JavaScript, Python, PHP, VB.NET, C#, SQL"),
    ("Frontend", "React, Tailwind CSS, HTML5, CSS3, Framer Motion, Vite"),
    ("Backend", "Node.js, Laravel, REST APIs, server-side application development"),
    ("Databases", "MySQL, MongoDB, Microsoft SQL Server, database design"),
    ("Tools and Platforms", "Git, GitHub Actions, Docker, Postman, Vercel, VS Code"),
    ("AI and Desktop", "MediaPipe, OpenCV, PySide6, Windows Forms, .NET Framework"),
]

EXPERIENCE = [
    {
        "role": "Backend Developer (Intern promoted to Full-Time)",
        "company": "Excellence and Innovation | Philippines",
        "dates": "May 2023 - Present",
        "bullets": [
            "Develop and maintain server-side application logic, REST APIs, and database workflows for web-based business systems.",
            "Manage data storage, retrieval, validation, and processing using PHP, MySQL, and Microsoft SQL Server.",
            "Integrate backend services with React and Tailwind CSS interfaces in collaboration with frontend developers.",
            "Support enrollment, cooperative, and accounting applications from requirements analysis through implementation.",
        ],
    }
]

PROJECTS = [
    {
        "name": "Hand Gesture Recognition System | 2026",
        "bullets": [
            "Built a real-time desktop application for hand tracking, custom gesture recognition, mouse control, and keyboard automation.",
            "Separated webcam capture, AI inference, and rendering into a multithreaded pipeline for responsive interaction.",
        ],
        "tech": "Python, MediaPipe, OpenCV, PySide6, NumPy, JSON, Computer Vision",
    },
    {
        "name": "Python File Locker | 2026",
        "bullets": [
            "Developed a desktop security application for encrypting files, folders, images, videos, and documents.",
            "Implemented authenticated encryption, Argon2id password derivation, encrypted metadata, compression, and optional key-file protection.",
        ],
        "tech": "Python, PySide6, AES-256-GCM, ChaCha20-Poly1305, Argon2id, PyInstaller",
    },
    {
        "name": "Laundry Management System | Lead Programmer",
        "bullets": [
            "Led development of a Windows desktop system for customer records, transactions, service pricing, inventory, receipts, and audit logs.",
            "Integrated Microsoft SQL Server data workflows and SIM800C-based status notifications for laundry operations.",
        ],
        "tech": "VB.NET, .NET Framework, Windows Forms, ADO.NET, Microsoft SQL Server, SIM800C",
    },
]

EDUCATION = [
    "Bachelor of Science in Information System | Richwell Colleges Incorporated | 2019 - 2023",
    "ICT - Programming | Atec Technological College | 2017 - 2019",
]

CERTIFICATIONS = [
    "Web Development NC III | Technical Education and Skills Development Authority (TESDA) | 2026",
    "Tech Trends: Data Analytics, Intermediate Session | Department of Information and Communications Technology (DICT) | 2022",
    "Tech Trends: Blockchain, Intermediate Session | Department of Information and Communications Technology (DICT) | 2022",
]


def add_docx_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "1E3A5F")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_docx_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_docx_section(document, title):
    paragraph = document.add_paragraph()
    set_docx_spacing(paragraph, before=5, after=3)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(30, 58, 95)
    add_docx_bottom_border(paragraph)


def add_docx_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    set_docx_spacing(paragraph, after=1, line=1.0)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.add_run(text)


def generate_docx():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(0)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_spacing(name, after=1)
    name_run = name.add_run(NAME)
    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(19)
    name_run.font.color.rgb = RGBColor(15, 23, 42)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_spacing(title, after=2)
    title_run = title.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = RGBColor(37, 99, 235)

    for line in (CONTACT, LINKS):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_docx_spacing(paragraph, after=1)
        paragraph.add_run(line)

    add_docx_section(document, "Professional Summary")
    paragraph = document.add_paragraph(SUMMARY)
    set_docx_spacing(paragraph, after=1, line=1.03)

    add_docx_section(document, "Technical Skills")
    for label, values in SKILLS:
        paragraph = document.add_paragraph()
        set_docx_spacing(paragraph, after=1)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(values)

    add_docx_section(document, "Professional Experience")
    for job in EXPERIENCE:
        paragraph = document.add_paragraph()
        set_docx_spacing(paragraph, after=0)
        role = paragraph.add_run(job["role"])
        role.bold = True
        paragraph.add_run(f" | {job['dates']}")
        company = document.add_paragraph(job["company"])
        set_docx_spacing(company, after=1)
        company.runs[0].italic = True
        for bullet in job["bullets"]:
            add_docx_bullet(document, bullet)

    add_docx_section(document, "Selected Projects")
    for project in PROJECTS:
        paragraph = document.add_paragraph()
        set_docx_spacing(paragraph, before=1, after=0)
        run = paragraph.add_run(project["name"])
        run.bold = True
        for bullet in project["bullets"]:
            add_docx_bullet(document, bullet)
        technologies = document.add_paragraph()
        set_docx_spacing(technologies, after=1)
        label = technologies.add_run("Technologies: ")
        label.bold = True
        technologies.add_run(project["tech"])

    add_docx_section(document, "Education")
    for item in EDUCATION:
        paragraph = document.add_paragraph(item)
        set_docx_spacing(paragraph, after=1)

    add_docx_section(document, "Certifications")
    for item in CERTIFICATIONS:
        paragraph = document.add_paragraph(item)
        set_docx_spacing(paragraph, after=1)

    document.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "ResumeName",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=21,
            leading=23,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#0F172A"),
            spaceAfter=1,
        ),
        "title": ParagraphStyle(
            "ResumeTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=12.5,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#2563EB"),
            spaceAfter=2,
        ),
        "contact": ParagraphStyle(
            "ResumeContact",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=10.5,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#334155"),
            spaceAfter=1,
        ),
        "section": ParagraphStyle(
            "ResumeSection",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=12.5,
            textColor=colors.HexColor("#1E3A5F"),
            spaceBefore=6,
            spaceAfter=3,
            borderWidth=0,
            borderPadding=0,
        ),
        "body": ParagraphStyle(
            "ResumeBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=11.2,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=1,
        ),
        "entry": ParagraphStyle(
            "ResumeEntry",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.6,
            leading=11.3,
            textColor=colors.HexColor("#111827"),
            spaceAfter=1,
        ),
        "subentry": ParagraphStyle(
            "ResumeSubentry",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=9,
            leading=10.6,
            textColor=colors.HexColor("#475569"),
            spaceAfter=1,
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.1,
            leading=11,
            leftIndent=11,
            firstLineIndent=-7,
            bulletIndent=2,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=0.5,
        ),
    }


def section_heading(text, styles):
    return [
        Spacer(1, 2),
        Paragraph(text.upper(), styles["section"]),
        Spacer(1, 0.5),
    ]


def generate_pdf():
    styles = pdf_styles()
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=0.48 * inch,
        leftMargin=0.48 * inch,
        topMargin=0.38 * inch,
        bottomMargin=0.38 * inch,
        title=f"{NAME} - ATS Resume",
        author="Shimron M. Guray",
        subject="Programming Specialist and Full-Stack Developer Resume",
    )

    story = [
        Paragraph(NAME, styles["name"]),
        Paragraph(TITLE, styles["title"]),
        Paragraph(CONTACT, styles["contact"]),
        Paragraph(LINKS, styles["contact"]),
    ]

    story += section_heading("Professional Summary", styles)
    story.append(Paragraph(SUMMARY, styles["body"]))

    story += section_heading("Technical Skills", styles)
    for label, values in SKILLS:
        story.append(Paragraph(f"<b>{label}:</b> {values}", styles["body"]))

    story += section_heading("Professional Experience", styles)
    for job in EXPERIENCE:
        content = [
            Paragraph(f"{job['role']} | {job['dates']}", styles["entry"]),
            Paragraph(job["company"], styles["subentry"]),
        ]
        content.extend(
            Paragraph(f"- {bullet}", styles["bullet"]) for bullet in job["bullets"]
        )
        story.append(KeepTogether(content))

    story += section_heading("Selected Projects", styles)
    for project in PROJECTS:
        content = [Paragraph(project["name"], styles["entry"])]
        content.extend(
            Paragraph(f"- {bullet}", styles["bullet"])
            for bullet in project["bullets"]
        )
        content.append(
            Paragraph(f"<b>Technologies:</b> {project['tech']}", styles["body"])
        )
        content.append(Spacer(1, 1.5))
        story.append(KeepTogether(content))

    story += section_heading("Education", styles)
    for item in EDUCATION:
        story.append(Paragraph(item, styles["body"]))

    story += section_heading("Certifications", styles)
    for item in CERTIFICATIONS:
        story.append(Paragraph(item, styles["body"]))

    document.build(story)


def main():
    PUBLIC.mkdir(parents=True, exist_ok=True)
    if PDF_PATH.exists() and not ORIGINAL_PATH.exists():
        shutil.copy2(PDF_PATH, ORIGINAL_PATH)
    generate_docx()
    generate_pdf()
    print(f"Generated {PDF_PATH}")
    print(f"Generated {DOCX_PATH}")
    print(f"Preserved original at {ORIGINAL_PATH}")


if __name__ == "__main__":
    main()
